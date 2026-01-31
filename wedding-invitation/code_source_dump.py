import os
from docx import Document
from docx.shared import Pt


def extract_all_files_to_word(folder_path: str, output_docx: str):
    print(f"[DEBUG] Starting extraction from folder: {folder_path}")
    print("[DEBUG] Mode: ALL FILE EXTENSIONS")

    if not os.path.isdir(folder_path):
        print("[ERROR] Provided path is not a directory")
        return

    document = Document()
    document.add_heading("Source Code Dump (All Extensions)", level=1)

    file_count = 0
    skipped_files = 0

    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            ext = os.path.splitext(file)[1].lower()

            print(f"[DEBUG] Processing file: {file_path}")

            document.add_page_break()
            document.add_heading(
                f"File: {file_path} ({ext if ext else 'no extension'})",
                level=2
            )

            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                print(f"[WARN] UTF-8 failed, retrying latin-1: {file_path}")
                try:
                    with open(file_path, "r", encoding="latin-1") as f:
                        content = f.read()
                except Exception as e:
                    skipped_files += 1
                    error_msg = f"[SKIPPED] Could not read file: {e}"
                    print(error_msg)
                    document.add_paragraph(error_msg)
                    continue
            except Exception as e:
                skipped_files += 1
                error_msg = f"[SKIPPED] Could not read file: {e}"
                print(error_msg)
                document.add_paragraph(error_msg)
                continue

            para = document.add_paragraph()
            run = para.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

            file_count += 1

    document.save(output_docx)

    print("[DEBUG] Extraction complete")
    print(f"[DEBUG] Total files added: {file_count}")
    print(f"[DEBUG] Total files skipped (binary/unreadable): {skipped_files}")
    print(f"[DEBUG] Word document saved as: {output_docx}")


# ===============================
# USAGE
# ===============================
if __name__ == "__main__":
    SOURCE_FOLDER = r"C:\Users\Ridhv\Desktop\startup\money\wedding-invitation"
    OUTPUT_FILE = "source_code_dump.docx"

    extract_all_files_to_word(SOURCE_FOLDER, OUTPUT_FILE)